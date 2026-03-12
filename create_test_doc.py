from docx import Document
import os

if not os.path.exists('documents'):
    os.makedirs('documents')

doc = Document()
doc.add_heading('Space Exploration Summary', 0)

doc.add_paragraph(
    "Space exploration is the use of astronomy and space technology to explore outer space. "
    "While the exploration of space is carried out mainly by astronomers with telescopes, "
    "its physical exploration though is conducted both by unmanned robotic space probes and human spaceflight."
)

doc.add_heading('The Apollo Program', level=1)
doc.add_paragraph(
    "The Apollo program was the third United States human spaceflight program carried out by the National Aeronautics and Space Administration (NASA), "
    "which succeeded in preparing and landing the first humans on the Moon from 1968 to 1972. "
    "It was first conceived during Dwight D. Eisenhower's administration as a three-man spacecraft."
)

doc.save('documents/space_exploration.docx')
print("Successfully created documents/space_exploration.docx")
